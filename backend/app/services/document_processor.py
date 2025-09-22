import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
from datetime import datetime
from loguru import logger

import pypdf
from docx import Document
from PIL import Image
import pytesseract
import markdown
from bs4 import BeautifulSoup
import chardet

from app.core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = settings.ALLOWED_EXTENSIONS
        self.max_file_size = settings.MAX_FILE_SIZE

    async def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        try:
            logger.info(f"Processing document: {file_path} of type: {file_type}")

            if file_type == "pdf":
                content = await self._process_pdf(file_path)
            elif file_type == "docx":
                content = await self._process_docx(file_path)
            elif file_type in ["txt", "md"]:
                content = await self._process_text(file_path)
            elif file_type in ["jpg", "jpeg", "png"]:
                content = await self._process_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            file_hash = self._calculate_file_hash(file_path)
            metadata = {
                "file_path": file_path,
                "filename": os.path.basename(file_path),
                "file_type": file_type,
                "file_size": os.path.getsize(file_path),
                "processed_at": datetime.utcnow().isoformat(),
                "file_hash": file_hash,
                "char_count": len(content),
                "word_count": len(content.split())
            }

            return {
                "content": content,
                "metadata": metadata,
                "chunks": self._chunk_text(content)
            }

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise

    async def _process_pdf(self, file_path: str) -> str:
        text_content = []

        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_content.append(f"Page {page_num + 1}:\n{text}")

        return "\n\n".join(text_content)

    async def _process_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_content.append(row_text)

        return "\n\n".join(text_content)

    async def _process_text(self, file_path: str) -> str:
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'

        with open(file_path, 'r', encoding=encoding) as file:
            content = file.read()

        if file_path.endswith('.md'):
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()

        return content

    async def _process_image(self, file_path: str) -> str:
        try:
            image = Image.open(file_path)

            if image.mode != 'RGB':
                image = image.convert('RGB')

            text = pytesseract.image_to_string(image)

            metadata_text = f"Image: {os.path.basename(file_path)}\n"
            metadata_text += f"Dimensions: {image.width}x{image.height}\n"
            metadata_text += f"Format: {image.format}\n\n"

            return metadata_text + "Extracted Text:\n" + text
        except Exception as e:
            logger.warning(f"OCR failed for {file_path}: {str(e)}")
            return f"Image file: {os.path.basename(file_path)} (OCR not available)"

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        chunks = []
        words = text.split()

        if len(words) <= chunk_size:
            return [{
                "content": text,
                "chunk_id": 0,
                "start_index": 0,
                "end_index": len(text)
            }]

        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)

            chunks.append({
                "content": chunk_text,
                "chunk_id": len(chunks),
                "start_index": i,
                "end_index": min(i + chunk_size, len(words))
            })

        return chunks

    def _calculate_file_hash(self, file_path: str) -> str:
        hash_sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def validate_file(self, file_path: str) -> bool:
        if not os.path.exists(file_path):
            return False

        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            return False

        file_extension = Path(file_path).suffix.lower().lstrip('.')
        return file_extension in self.supported_extensions